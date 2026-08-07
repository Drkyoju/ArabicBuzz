#!/usr/bin/env python3
"""Local PDF↔DOCX helper for ArabicBuzz agents (Mac / storage:sync).

Quality order (do not invent a worse path):
  1) LibreOffice soffice --headless (if installed) — only when Unicode looks healthy
  2) pdf2docx (layout) — WARN if Arabic ToUnicode looks broken
  3) Visual page-image DOCX (pixel layout match; not text-editable)

Never use pdf-lib stamping for Arabic body text.
Prefer Google Drive convert on the live site when linked (OCR + editable).
"""
from __future__ import annotations

import argparse
import io
import re
import shutil
import subprocess
import sys
from pathlib import Path


BROKEN_AR_HINTS = (
    "املادة",
    "الالئحة",
    "االسم",
    "األساسية",
    "واألهداف",
)


def arabic_text_looks_broken(sample: str) -> bool:
    hits = sum(1 for h in BROKEN_AR_HINTS if h in sample)
    if hits >= 2:
        return True
    bad = len(re.findall(r"اال|امل[^ا]", sample))
    good = len(re.findall(r"ال[اأإ]|الم", sample))
    return bad > good * 1.5 and bad > 8


def find_soffice() -> str | None:
    candidates = [
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/local/bin/soffice",
        "/opt/homebrew/bin/soffice",
    ]
    for c in candidates:
        if c and Path(c).exists():
            return c
    return None


def convert_soffice(src: Path, out_dir: Path, to_ext: str) -> Path:
    soffice = find_soffice()
    if not soffice:
        raise RuntimeError("soffice غير متوفر")
    out_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        to_ext,
        "--outdir",
        str(out_dir),
        str(src),
    ]
    subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=300)
    out = out_dir / f"{src.stem}.{to_ext}"
    if not out.exists():
        raise RuntimeError(f"LibreOffice لم يُنتج {out.name}")
    return out


def convert_pdf2docx(src: Path, dest: Path) -> Path:
    from pdf2docx import Converter

    dest.parent.mkdir(parents=True, exist_ok=True)
    cv = Converter(str(src))
    try:
        cv.convert(str(dest))
    finally:
        cv.close()
    return dest


def convert_visual_docx(src: Path, dest: Path, dpi: int = 180) -> Path:
    """One full-bleed page image per Word page — layout matches PDF visually 100%.

    Not text-editable. Use when ToUnicode is broken (Sakkal Majalla etc.).
    Uses JPEG under ~4MB target when page count is high (cloud chat upload).
    """
    import fitz
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Emu, Pt, Twips

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)

    pdf = fitz.open(src)
    page_count = pdf.page_count
    # Prefer JPEG for long docs so cloud (4MB) + room chat stay possible
    use_jpeg = page_count >= 15 or dpi <= 130
    jpeg_quality = 68 if page_count >= 30 else 78
    mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    PT_TO_TWIPS = 20

    for i, page in enumerate(pdf):
        pw_pt = float(page.rect.width)
        ph_pt = float(page.rect.height)
        page_w = Twips(int(round(pw_pt * PT_TO_TWIPS)))
        page_h = Twips(int(round(ph_pt * PT_TO_TWIPS)))

        if i == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section()

        section.page_width = page_w
        section.page_height = page_h
        section.orientation = (
            WD_ORIENT.LANDSCAPE if pw_pt > ph_pt else WD_ORIENT.PORTRAIT
        )
        section.top_margin = Twips(0)
        section.bottom_margin = Twips(0)
        section.left_margin = Twips(0)
        section.right_margin = Twips(0)
        section.header_distance = Twips(0)
        section.footer_distance = Twips(0)

        pix = page.get_pixmap(matrix=mat, alpha=False)
        if use_jpeg:
            try:
                from PIL import Image

                im = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                img = io.BytesIO()
                im.save(img, format="JPEG", quality=jpeg_quality, optimize=True)
                img.seek(0)
            except Exception:
                img = io.BytesIO(pix.tobytes("jpeg", jpg_quality=jpeg_quality))
        else:
            img = io.BytesIO(pix.tobytes("png"))

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)

        run = para.add_run()
        width_emu = Emu(int(round(pw_pt / 72.0 * 914400)))
        run.add_picture(img, width=width_emu)

    try:
        doc.core_properties.comments = (
            "نسخة مرئية مطابقة لتخطيط PDF (صورة لكل صفحة). "
            "طبقة Unicode في المصدر قد تكون معطوبة — للتحرير النصي اربط Google Drive."
        )
        doc.core_properties.subject = "visual-pdf-match"
    except Exception:
        pass

    doc.save(dest)
    return dest


def sample_pdf_text(src: Path, max_pages: int = 3) -> str:
    import fitz

    pdf = fitz.open(src)
    parts = []
    for i, page in enumerate(pdf):
        if i >= max_pages:
            break
        parts.append(page.get_text())
    return "\n".join(parts)


def main() -> int:
    ap = argparse.ArgumentParser(description="ArabicBuzz PDF↔DOCX local convert")
    ap.add_argument("input", type=Path)
    ap.add_argument("-o", "--output", type=Path, required=True)
    ap.add_argument(
        "--mode",
        choices=["auto", "soffice", "pdf2docx", "visual"],
        default="auto",
    )
    ap.add_argument("--dpi", type=int, default=180, help="DPI for visual mode")
    args = ap.parse_args()
    src = args.input.expanduser().resolve()
    dest = args.output.expanduser().resolve()
    if not src.exists():
        print(f"missing: {src}", file=sys.stderr)
        return 2

    to_docx = src.suffix.lower() == ".pdf" and dest.suffix.lower() == ".docx"
    to_pdf = src.suffix.lower() in {".docx", ".doc"} and dest.suffix.lower() == ".pdf"

    mode = args.mode
    warnings: list[str] = []

    if to_docx:
        sample = sample_pdf_text(src)
        broken = arabic_text_looks_broken(sample)
        if broken:
            warnings.append(
                "طبقة Unicode في PDF تبدو معطوبة للعربية — تجنّب المسار النصّي؛ "
                "فضّل Google Drive / visual / OCR."
            )

        if mode == "auto":
            if find_soffice() and not broken:
                mode = "soffice"
            elif broken:
                mode = "visual"
            else:
                mode = "pdf2docx"

        try:
            if mode == "soffice":
                tmp = convert_soffice(src, dest.parent, "docx")
                if tmp != dest:
                    shutil.move(str(tmp), str(dest))
            elif mode == "pdf2docx":
                convert_pdf2docx(src, dest)
            else:
                convert_visual_docx(src, dest, dpi=args.dpi)
        except Exception as e:
            if mode != "visual":
                warnings.append(f"{mode} فشل ({e}); الانتقال للنسخة المرئية")
                convert_visual_docx(src, dest, dpi=args.dpi)
                mode = "visual"
            else:
                raise
    elif to_pdf:
        if mode in ("auto", "soffice"):
            tmp = convert_soffice(src, dest.parent, "pdf")
            if tmp != dest:
                shutil.move(str(tmp), str(dest))
            mode = "soffice"
        else:
            print("Word→PDF المحلي يحتاج LibreOffice (soffice)", file=sys.stderr)
            return 3
    else:
        print("زوج الصيغ غير مدعوم محلياً", file=sys.stderr)
        return 4

    print(f"ok engine={mode} out={dest}")
    for w in warnings:
        print(f"WARN: {w}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
